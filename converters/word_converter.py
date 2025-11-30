"""
Word Converter
Handles DOCX conversions
"""

from pathlib import Path
from docx import Document
import json
from .base_converter import BaseConverter
from utils.errors import ConversionError


class WordConverter(BaseConverter):
    """Convert DOCX to other formats"""
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['docx']
        self.supported_output_formats = ['txt', 'json']
    
    def convert(
        self,
        input_file: str | Path,
        output_file: str | Path,
        **kwargs
    ) -> Path:
        """
        Convert DOCX file
        
        Args:
            input_file: Path to DOCX file
            output_file: Path to output file
            **kwargs: Additional options
            
        Returns:
            Path to converted file
        """
        input_file = Path(input_file)
        output_file = Path(output_file)
        
        self._ensure_output_directory(output_file)
        
        output_format = output_file.suffix.lower().replace('.', '')
        
        # Route to appropriate converter
        if output_format == 'txt':
            return self._docx_to_txt(input_file, output_file, **kwargs)
        elif output_format == 'json':
            return self._docx_to_json(input_file, output_file, **kwargs)
        else:
            raise ConversionError(f"Unsupported output format: {output_format}")
    
    def _docx_to_txt(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert DOCX to TXT"""
        try:
            doc = Document(input_file)
            
            content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                content.append("\n--- Table ---")
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content.append(row_text)
                content.append("--- End Table ---\n")
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))
            
            self._log_info(f"Converted DOCX to TXT: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"DOCX to TXT conversion failed: {e}")
    
    def _docx_to_json(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert DOCX to JSON"""
        try:
            doc = Document(input_file)
            
            doc_data = {
                'paragraphs': [],
                'tables': []
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    para_data = {
                        'text': para.text,
                        'style': para.style.name
                    }
                    doc_data['paragraphs'].append(para_data)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'table_id': table_idx + 1,
                    'rows': []
                }
                
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data['rows'].append(row_data)
                
                doc_data['tables'].append(table_data)
            
            # Write to JSON
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, indent=2, ensure_ascii=False)
            
            self._log_info(f"Converted DOCX to JSON: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"DOCX to JSON conversion failed: {e}")


class ToWordConverter(BaseConverter):
    """Convert other formats to DOCX"""
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['txt', 'json']
        self.supported_output_formats = ['docx']
    
    def convert(
        self,
        input_file: str | Path,
        output_file: str | Path,
        **kwargs
    ) -> Path:
        """Convert to DOCX"""
        input_file = Path(input_file)
        output_file = Path(output_file)
        
        self._ensure_output_directory(output_file)
        
        input_format = input_file.suffix.lower().replace('.', '')
        
        # Route to appropriate converter
        if input_format == 'txt':
            return self._txt_to_docx(input_file, output_file, **kwargs)
        elif input_format == 'json':
            return self._json_to_docx(input_file, output_file, **kwargs)
        else:
            raise ConversionError(f"Unsupported input format: {input_format}")
    
    def _txt_to_docx(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert TXT to DOCX"""
        try:
            doc = Document()
            
            # Read text file
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split into paragraphs
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Check if it looks like a heading
                    lines = para.split('\n')
                    for line in lines:
                        if line.strip():
                            # Simple heuristic: short lines might be headings
                            if len(line) < 50 and line.isupper():
                                doc.add_heading(line.strip(), level=1)
                            elif len(line) < 50 and line.strip().endswith(':'):
                                doc.add_heading(line.strip(), level=2)
                            else:
                                doc.add_paragraph(line.strip())
            
            doc.save(output_file)
            self._log_info(f"Converted TXT to DOCX: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"TXT to DOCX conversion failed: {e}")
    
    def _json_to_docx(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert JSON to DOCX"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            doc = Document()
            
            # Handle different JSON structures
            if isinstance(data, dict):
                # Check if it has document structure
                if 'paragraphs' in data:
                    # Structured document
                    for para_data in data.get('paragraphs', []):
                        if isinstance(para_data, dict):
                            text = para_data.get('text', '')
                            style = para_data.get('style', 'Normal')
                            
                            if 'Heading' in style:
                                level = int(style.replace('Heading', '').strip() or '1')
                                doc.add_heading(text, level=level)
                            else:
                                doc.add_paragraph(text)
                        else:
                            doc.add_paragraph(str(para_data))
                    
                    # Add tables if present
                    for table_data in data.get('tables', []):
                        rows = table_data.get('rows', [])
                        if rows:
                            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                            table.style = 'Light Grid Accent 1'
                            
                            for i, row in enumerate(rows):
                                for j, cell in enumerate(row):
                                    table.rows[i].cells[j].text = str(cell)
                else:
                    # Generic dict - format as key-value pairs
                    for key, value in data.items():
                        doc.add_heading(str(key), level=2)
                        doc.add_paragraph(str(value))
            
            elif isinstance(data, list):
                # List of items
                for item in data:
                    if isinstance(item, dict):
                        for key, value in item.items():
                            doc.add_paragraph(f"{key}: {value}")
                    else:
                        doc.add_paragraph(str(item))
            else:
                # Simple value
                doc.add_paragraph(str(data))
            
            doc.save(output_file)
            self._log_info(f"Converted JSON to DOCX: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"JSON to DOCX conversion failed: {e}")
