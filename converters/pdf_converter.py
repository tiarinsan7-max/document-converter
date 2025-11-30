"""
PDF Converter
Handles all PDF-related conversions
"""

from pathlib import Path
from typing import Optional
import pdfplumber
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import pandas as pd
from docx import Document
import json
from .base_converter import BaseConverter
from utils.errors import ConversionError


class PDFConverter(BaseConverter):
    """Convert PDF to other formats and vice versa"""
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['pdf']
        self.supported_output_formats = ['txt', 'docx', 'csv', 'xlsx', 'json']
    
    def convert(
        self,
        input_file: str | Path,
        output_file: str | Path,
        **kwargs
    ) -> Path:
        """
        Convert PDF to target format
        
        Args:
            input_file: Path to PDF file
            output_file: Path to output file
            **kwargs: Additional options
            
        Returns:
            Path to converted file
        """
        input_file = Path(input_file)
        output_file = Path(output_file)
        
        self._ensure_output_directory(output_file)
        
        output_format = output_file.suffix.lower().replace('.', '')
        
        # Route to appropriate converter
        if output_format == 'txt':
            return self._pdf_to_txt(input_file, output_file, **kwargs)
        elif output_format == 'docx':
            return self._pdf_to_docx(input_file, output_file, **kwargs)
        elif output_format == 'csv':
            return self._pdf_to_csv(input_file, output_file, **kwargs)
        elif output_format == 'xlsx':
            return self._pdf_to_xlsx(input_file, output_file, **kwargs)
        elif output_format == 'json':
            return self._pdf_to_json(input_file, output_file, **kwargs)
        else:
            raise ConversionError(f"Unsupported output format: {output_format}")
    
    def _pdf_to_txt(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert PDF to TXT"""
        try:
            text_content = []
            
            with pdfplumber.open(input_file) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"--- Page {page_num} ---\n")
                        text_content.append(text)
                        text_content.append("\n\n")
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(''.join(text_content))
            
            self._log_info(f"Converted PDF to TXT: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"PDF to TXT conversion failed: {e}")
    
    def _pdf_to_docx(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert PDF to DOCX"""
        try:
            doc = Document()
            
            with pdfplumber.open(input_file) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Add page heading
                    doc.add_heading(f'Page {page_num}', level=2)
                    
                    # Extract text
                    text = page.extract_text()
                    if text:
                        # Split into paragraphs
                        paragraphs = text.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para.strip())
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table_data in tables:
                            if table_data:
                                # Add table to document
                                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                                table.style = 'Light Grid Accent 1'
                                
                                for i, row in enumerate(table_data):
                                    for j, cell in enumerate(row):
                                        if cell:
                                            table.rows[i].cells[j].text = str(cell)
                    
                    # Page break (except for last page)
                    if page_num < len(pdf.pages):
                        doc.add_page_break()
            
            doc.save(output_file)
            self._log_info(f"Converted PDF to DOCX: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"PDF to DOCX conversion failed: {e}")
    
    def _pdf_to_csv(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert PDF tables to CSV"""
        try:
            all_tables = []
            
            with pdfplumber.open(input_file) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table:
                                all_tables.extend(table)
            
            if not all_tables:
                # If no tables, extract text as single column
                text_content = []
                with pdfplumber.open(input_file) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            lines = text.split('\n')
                            text_content.extend([[line] for line in lines if line.strip()])
                
                df = pd.DataFrame(text_content, columns=['Content'])
            else:
                df = pd.DataFrame(all_tables)
            
            df.to_csv(output_file, index=False, encoding='utf-8')
            self._log_info(f"Converted PDF to CSV: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"PDF to CSV conversion failed: {e}")
    
    def _pdf_to_xlsx(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert PDF to XLSX"""
        try:
            all_tables = []
            
            with pdfplumber.open(input_file) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, 1):
                            if table:
                                df = pd.DataFrame(table)
                                all_tables.append((f"Page{page_num}_Table{table_num}", df))
            
            if not all_tables:
                # Extract text as single sheet
                text_content = []
                with pdfplumber.open(input_file) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            lines = text.split('\n')
                            text_content.extend([[line] for line in lines if line.strip()])
                
                df = pd.DataFrame(text_content, columns=['Content'])
                all_tables.append(("Content", df))
            
            # Write to Excel
            with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                for sheet_name, df in all_tables:
                    # Limit sheet name to 31 characters
                    sheet_name = sheet_name[:31]
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            self._log_info(f"Converted PDF to XLSX: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"PDF to XLSX conversion failed: {e}")
    
    def _pdf_to_json(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert PDF to JSON"""
        try:
            pdf_data = {
                'metadata': {},
                'pages': []
            }
            
            with pdfplumber.open(input_file) as pdf:
                # Extract metadata
                if pdf.metadata:
                    pdf_data['metadata'] = {
                        k: str(v) for k, v in pdf.metadata.items() if v
                    }
                
                # Extract pages
                for page_num, page in enumerate(pdf.pages, 1):
                    page_data = {
                        'page_number': page_num,
                        'text': page.extract_text() or '',
                        'tables': []
                    }
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table:
                                page_data['tables'].append(table)
                    
                    pdf_data['pages'].append(page_data)
            
            # Write to JSON
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(pdf_data, f, indent=2, ensure_ascii=False)
            
            self._log_info(f"Converted PDF to JSON: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"PDF to JSON conversion failed: {e}")


class ToPDFConverter(BaseConverter):
    """Convert other formats to PDF"""
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['txt', 'csv', 'xlsx', 'json', 'docx']
        self.supported_output_formats = ['pdf']
    
    def convert(
        self,
        input_file: str | Path,
        output_file: str | Path,
        **kwargs
    ) -> Path:
        """Convert to PDF"""
        input_file = Path(input_file)
        output_file = Path(output_file)
        
        self._ensure_output_directory(output_file)
        
        input_format = input_file.suffix.lower().replace('.', '')
        
        # Route to appropriate converter
        if input_format == 'txt':
            return self._txt_to_pdf(input_file, output_file, **kwargs)
        elif input_format == 'csv':
            return self._csv_to_pdf(input_file, output_file, **kwargs)
        elif input_format == 'xlsx':
            return self._xlsx_to_pdf(input_file, output_file, **kwargs)
        elif input_format == 'json':
            return self._json_to_pdf(input_file, output_file, **kwargs)
        elif input_format == 'docx':
            return self._docx_to_pdf(input_file, output_file, **kwargs)
        else:
            raise ConversionError(f"Unsupported input format: {input_format}")
    
    def _txt_to_pdf(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert TXT to PDF"""
        try:
            doc = SimpleDocTemplate(str(output_file), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Read text file
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split into paragraphs
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Handle long lines
                    lines = para.split('\n')
                    for line in lines:
                        if line.strip():
                            p = Paragraph(line.strip(), styles['Normal'])
                            story.append(p)
                            story.append(Spacer(1, 0.1 * inch))
            
            doc.build(story)
            self._log_info(f"Converted TXT to PDF: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"TXT to PDF conversion failed: {e}")
    
    def _csv_to_pdf(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert CSV to PDF"""
        try:
            df = pd.read_csv(input_file)
            
            doc = SimpleDocTemplate(str(output_file), pagesize=A4)
            story = []
            
            # Convert DataFrame to list for table
            data = [df.columns.tolist()] + df.values.tolist()
            
            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            doc.build(story)
            
            self._log_info(f"Converted CSV to PDF: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"CSV to PDF conversion failed: {e}")
    
    def _xlsx_to_pdf(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert XLSX to PDF"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(input_file)
            
            doc = SimpleDocTemplate(str(output_file), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(input_file, sheet_name=sheet_name)
                
                # Add sheet title
                story.append(Paragraph(f"<b>{sheet_name}</b>", styles['Heading1']))
                story.append(Spacer(1, 0.2 * inch))
                
                # Convert to table
                data = [df.columns.tolist()] + df.values.tolist()
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
                story.append(Spacer(1, 0.3 * inch))
            
            doc.build(story)
            self._log_info(f"Converted XLSX to PDF: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"XLSX to PDF conversion failed: {e}")
    
    def _json_to_pdf(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert JSON to PDF"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            doc = SimpleDocTemplate(str(output_file), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Format JSON as readable text
            json_str = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Split into lines
            lines = json_str.split('\n')
            
            for line in lines:
                if line.strip():
                    p = Paragraph(line.replace(' ', '&nbsp;'), styles['Code'])
                    story.append(p)
            
            doc.build(story)
            self._log_info(f"Converted JSON to PDF: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"JSON to PDF conversion failed: {e}")
    
    def _docx_to_pdf(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert DOCX to PDF"""
        try:
            # Read DOCX
            doc_input = Document(input_file)
            
            # Create PDF
            doc_output = SimpleDocTemplate(str(output_file), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Extract paragraphs
            for para in doc_input.paragraphs:
                if para.text.strip():
                    # Determine style based on paragraph style
                    if para.style.name.startswith('Heading'):
                        style = styles['Heading1']
                    else:
                        style = styles['Normal']
                    
                    p = Paragraph(para.text, style)
                    story.append(p)
                    story.append(Spacer(1, 0.1 * inch))
            
            # Extract tables
            for table in doc_input.tables:
                data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    data.append(row_data)
                
                if data:
                    pdf_table = Table(data)
                    pdf_table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 0.2 * inch))
            
            doc_output.build(story)
            self._log_info(f"Converted DOCX to PDF: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"DOCX to PDF conversion failed: {e}")
