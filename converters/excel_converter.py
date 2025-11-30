"""
Excel Converter
Handles CSV and XLSX conversions
"""

from pathlib import Path
import pandas as pd
import json
from docx import Document
from docx.shared import Inches
from .base_converter import BaseConverter
from utils.errors import ConversionError


class ExcelConverter(BaseConverter):
    """Convert between CSV, XLSX, and other formats"""
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['csv', 'xlsx']
        self.supported_output_formats = ['csv', 'xlsx', 'json', 'txt', 'docx']
    
    def convert(
        self,
        input_file: str | Path,
        output_file: str | Path,
        **kwargs
    ) -> Path:
        """
        Convert Excel/CSV file
        
        Args:
            input_file: Path to input file
            output_file: Path to output file
            **kwargs: Additional options
            
        Returns:
            Path to converted file
        """
        input_file = Path(input_file)
        output_file = Path(output_file)
        
        self._ensure_output_directory(output_file)
        
        input_format = input_file.suffix.lower().replace('.', '')
        output_format = output_file.suffix.lower().replace('.', '')
        
        # Route to appropriate converter
        if input_format == 'csv' and output_format == 'xlsx':
            return self._csv_to_xlsx(input_file, output_file, **kwargs)
        elif input_format == 'xlsx' and output_format == 'csv':
            return self._xlsx_to_csv(input_file, output_file, **kwargs)
        elif input_format in ['csv', 'xlsx'] and output_format == 'json':
            return self._excel_to_json(input_file, output_file, **kwargs)
        elif input_format in ['csv', 'xlsx'] and output_format == 'txt':
            return self._excel_to_txt(input_file, output_file, **kwargs)
        elif input_format in ['csv', 'xlsx'] and output_format == 'docx':
            return self._excel_to_docx(input_file, output_file, **kwargs)
        else:
            raise ConversionError(
                f"Unsupported conversion: {input_format} to {output_format}"
            )
    
    def _csv_to_xlsx(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert CSV to XLSX"""
        try:
            # Read CSV
            df = pd.read_csv(input_file, encoding='utf-8')
            
            # Write to Excel
            with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Sheet1', index=False)
                
                # Auto-adjust column widths
                worksheet = writer.sheets['Sheet1']
                for idx, col in enumerate(df.columns):
                    max_length = max(
                        df[col].astype(str).apply(len).max(),
                        len(str(col))
                    )
                    worksheet.column_dimensions[chr(65 + idx)].width = min(max_length + 2, 50)
            
            self._log_info(f"Converted CSV to XLSX: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"CSV to XLSX conversion failed: {e}")
    
    def _xlsx_to_csv(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert XLSX to CSV"""
        try:
            # Read first sheet by default
            sheet_name = kwargs.get('sheet_name', 0)
            df = pd.read_excel(input_file, sheet_name=sheet_name)
            
            # Write to CSV
            df.to_csv(output_file, index=False, encoding='utf-8')
            
            self._log_info(f"Converted XLSX to CSV: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"XLSX to CSV conversion failed: {e}")
    
    def _excel_to_json(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert CSV/XLSX to JSON"""
        try:
            input_format = input_file.suffix.lower().replace('.', '')
            
            if input_format == 'csv':
                df = pd.read_csv(input_file, encoding='utf-8')
                data = {
                    'data': df.to_dict(orient='records')
                }
            else:  # xlsx
                excel_file = pd.ExcelFile(input_file)
                data = {}
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(input_file, sheet_name=sheet_name)
                    data[sheet_name] = df.to_dict(orient='records')
            
            # Write to JSON
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False, default=str)
            
            self._log_info(f"Converted {input_format.upper()} to JSON: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"Excel to JSON conversion failed: {e}")
    
    def _excel_to_txt(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert CSV/XLSX to TXT"""
        try:
            input_format = input_file.suffix.lower().replace('.', '')
            
            if input_format == 'csv':
                df = pd.read_csv(input_file, encoding='utf-8')
                content = df.to_string(index=False)
            else:  # xlsx
                excel_file = pd.ExcelFile(input_file)
                content_parts = []
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(input_file, sheet_name=sheet_name)
                    content_parts.append(f"=== {sheet_name} ===\n")
                    content_parts.append(df.to_string(index=False))
                    content_parts.append("\n\n")
                
                content = ''.join(content_parts)
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            
            self._log_info(f"Converted {input_format.upper()} to TXT: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"Excel to TXT conversion failed: {e}")
    
    def _excel_to_docx(self, input_file: Path, output_file: Path, **kwargs) -> Path:
        """Convert CSV/XLSX to DOCX"""
        try:
            input_format = input_file.suffix.lower().replace('.', '')
            doc = Document()
            
            if input_format == 'csv':
                df = pd.read_csv(input_file, encoding='utf-8')
                self._add_dataframe_to_doc(doc, df, "Data")
            else:  # xlsx
                excel_file = pd.ExcelFile(input_file)
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(input_file, sheet_name=sheet_name)
                    
                    # Add sheet heading
                    doc.add_heading(sheet_name, level=1)
                    
                    # Add table
                    self._add_dataframe_to_doc(doc, df)
                    
                    # Add page break (except for last sheet)
                    if sheet_name != excel_file.sheet_names[-1]:
                        doc.add_page_break()
            
            doc.save(output_file)
            self._log_info(f"Converted {input_format.upper()} to DOCX: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"Excel to DOCX conversion failed: {e}")
    
    def _add_dataframe_to_doc(self, doc: Document, df: pd.DataFrame, title: str = None):
        """Add DataFrame as table to Word document"""
        if title:
            doc.add_heading(title, level=2)
        
        # Create table
        table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        for i, column in enumerate(df.columns):
            table.rows[0].cells[i].text = str(column)
        
        # Add data
        for i, row in enumerate(df.itertuples(index=False), 1):
            for j, value in enumerate(row):
                table.rows[i].cells[j].text = str(value) if pd.notna(value) else ''


class FromExcelConverter(BaseConverter):
    """Convert from JSON/TXT/DOCX to CSV/XLSX"""
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['json', 'txt', 'docx']
        self.supported_output_formats = ['csv', 'xlsx']
    
    def convert(
        self,
        input_file: str | Path,
        output_file: str | Path,
        **kwargs
    ) -> Path:
        """Convert to CSV/XLSX"""
        input_file = Path(input_file)
        output_file = Path(output_file)
        
        self._ensure_output_directory(output_file)
        
        input_format = input_file.suffix.lower().replace('.', '')
        output_format = output_file.suffix.lower().replace('.', '')
        
        # Route to appropriate converter
        if input_format == 'json':
            return self._json_to_excel(input_file, output_file, output_format, **kwargs)
        elif input_format == 'txt':
            return self._txt_to_excel(input_file, output_file, output_format, **kwargs)
        elif input_format == 'docx':
            return self._docx_to_excel(input_file, output_file, output_format, **kwargs)
        else:
            raise ConversionError(f"Unsupported input format: {input_format}")
    
    def _json_to_excel(self, input_file: Path, output_file: Path, output_format: str, **kwargs) -> Path:
        """Convert JSON to CSV/XLSX"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Handle different JSON structures
            if isinstance(data, list):
                df = pd.DataFrame(data)
            elif isinstance(data, dict):
                # Check if it's a multi-sheet structure
                if all(isinstance(v, list) for v in data.values()):
                    # Multi-sheet
                    if output_format == 'xlsx':
                        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                            for sheet_name, sheet_data in data.items():
                                df = pd.DataFrame(sheet_data)
                                df.to_excel(writer, sheet_name=str(sheet_name)[:31], index=False)
                        
                        self._log_info(f"Converted JSON to XLSX: {output_file}")
                        return output_file
                    else:
                        # For CSV, use first sheet
                        first_key = list(data.keys())[0]
                        df = pd.DataFrame(data[first_key])
                else:
                    # Single object, convert to single-row DataFrame
                    df = pd.DataFrame([data])
            else:
                raise ConversionError("Unsupported JSON structure")
            
            # Write to file
            if output_format == 'csv':
                df.to_csv(output_file, index=False, encoding='utf-8')
            else:  # xlsx
                df.to_excel(output_file, index=False, engine='openpyxl')
            
            self._log_info(f"Converted JSON to {output_format.upper()}: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"JSON to Excel conversion failed: {e}")
    
    def _txt_to_excel(self, input_file: Path, output_file: Path, output_format: str, **kwargs) -> Path:
        """Convert TXT to CSV/XLSX"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            # Try to detect delimiter
            delimiter = kwargs.get('delimiter', ',')
            
            # Create DataFrame
            data = [line.strip().split(delimiter) for line in lines if line.strip()]
            df = pd.DataFrame(data)
            
            # Write to file
            if output_format == 'csv':
                df.to_csv(output_file, index=False, header=False, encoding='utf-8')
            else:  # xlsx
                df.to_excel(output_file, index=False, header=False, engine='openpyxl')
            
            self._log_info(f"Converted TXT to {output_format.upper()}: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"TXT to Excel conversion failed: {e}")
    
    def _docx_to_excel(self, input_file: Path, output_file: Path, output_format: str, **kwargs) -> Path:
        """Convert DOCX tables to CSV/XLSX"""
        try:
            doc = Document(input_file)
            
            # Extract all tables
            all_data = []
            
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    all_data.append(row_data)
            
            if not all_data:
                # No tables, extract text as single column
                all_data = [[para.text] for para in doc.paragraphs if para.text.strip()]
            
            df = pd.DataFrame(all_data)
            
            # Write to file
            if output_format == 'csv':
                df.to_csv(output_file, index=False, header=False, encoding='utf-8')
            else:  # xlsx
                df.to_excel(output_file, index=False, header=False, engine='openpyxl')
            
            self._log_info(f"Converted DOCX to {output_format.upper()}: {output_file}")
            return output_file
            
        except Exception as e:
            raise ConversionError(f"DOCX to Excel conversion failed: {e}")
